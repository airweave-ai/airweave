"""Word bongo implementation.

Creates, updates, and deletes test Word documents via the Microsoft Graph API.
"""

import asyncio
import io
import uuid
from typing import Any, Dict, List

import httpx
from monke.bongos._microsoft_graph_base import GRAPH, MicrosoftGraphBongo
from monke.generation.word import generate_documents_content

# Try to import python-docx for Word document creation
try:
    from docx import Document
    from docx.shared import Pt

    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False


class WordBongo(MicrosoftGraphBongo):
    """Bongo for Word that creates test entities for E2E testing.

    Key responsibilities:
    - Create test Word documents in OneDrive
    - Add content with verification tokens
    - Update documents to test incremental sync
    - Delete documents to test deletion detection
    - Clean up all test data
    """

    connector_type = "word"

    def __init__(self, credentials: Dict[str, Any], **kwargs):
        super().__init__(credentials, **kwargs)

        # Track created resources for cleanup
        self._test_documents: List[Dict[str, Any]] = []

        if not HAS_PYTHON_DOCX:
            raise ImportError(
                "python-docx is required for Word bongo. Install with: pip install python-docx"
            )

    async def create_entities(self) -> List[Dict[str, Any]]:
        """Create test Word documents in OneDrive."""
        self.logger.info(f"🥁 Creating {self.entity_count} Word test documents")
        out: List[Dict[str, Any]] = []

        # Generate tokens for each document
        tokens = [uuid.uuid4().hex[:8] for _ in range(self.entity_count)]

        # Generate document content
        test_name = f"TestDoc_{uuid.uuid4().hex[:8]}"
        filenames, document_content = await generate_documents_content(
            self.openai_model, tokens, test_name
        )

        self.logger.info(f"📄 Generated {len(document_content)} documents")

        async with httpx.AsyncClient(base_url=GRAPH, timeout=60) as client:
            for i, (filename, doc_content, token) in enumerate(
                zip(filenames, document_content, tokens)
            ):
                await self._pace()

                # Sanitize filename for OneDrive (remove illegal characters)
                safe_filename = self._sanitize_filename(filename)

                # Create Word document file
                doc_bytes = self._create_word_file(doc_content)

                # Upload to OneDrive
                self.logger.info(f"📤 Uploading Word document: {safe_filename}")
                upload_url = f"/me/drive/root:/{safe_filename}:/content"

                r = await client.put(
                    upload_url,
                    headers={
                        "Authorization": f"Bearer {self.access_token}",
                        "Content-Type": (
                            "application/vnd.openxmlformats-officedocument."
                            "wordprocessingml.document"
                        ),
                    },
                    content=doc_bytes,
                )

                if r.status_code not in (200, 201):
                    self.logger.error(f"Upload failed {r.status_code}: {r.text}")
                    r.raise_for_status()

                doc_file = r.json()
                doc_id = doc_file["id"]

                self.logger.info(f"✅ Uploaded document: {doc_id} - {filename}")

                ent = {
                    "type": "document",
                    "id": doc_id,
                    "filename": safe_filename,
                    "title": doc_content.title,
                    "token": token,
                    "expected_content": token,
                }
                out.append(ent)
                self._test_documents.append(ent)
                self.created_entities.append({"id": doc_id, "name": safe_filename})

                self.logger.info(f"📝 Document '{safe_filename}' created with token: {token}")

        self.logger.info(f"✅ Created {len(self._test_documents)} Word documents")
        return out

    def _create_word_file(self, doc_content: Any) -> bytes:
        """Create a Word document file with the given content.

        Args:
            doc_content: WordDocumentContent object

        Returns:
            Bytes of the Word document
        """
        doc = Document()

        # Add title
        title = doc.add_heading(doc_content.title, level=0)
        title.alignment = 1  # Center alignment

        # Add summary section
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(doc_content.summary)

        # Add main content
        doc.add_heading("Content", level=1)

        # Split content by paragraphs and add them
        paragraphs = doc_content.content.split("\n\n")
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph(para.strip())
                # Set font size for readability
                for run in p.runs:
                    run.font.size = Pt(11)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    async def update_entities(self) -> List[Dict[str, Any]]:
        """Update Word documents by appending new content with same tokens."""
        if not self._test_documents:
            return []

        self.logger.info(f"🥁 Updating {min(2, len(self._test_documents))} Word documents")
        updated = []

        async with httpx.AsyncClient(base_url=GRAPH, timeout=60) as client:
            for ent in self._test_documents[: min(2, len(self._test_documents))]:
                await self._pace()

                # Download current document
                download_url = f"/me/drive/items/{ent['id']}/content"
                r = await client.get(download_url, headers=self._hdrs())

                if r.status_code != 200:
                    self.logger.warning(
                        f"Failed to download document: {r.status_code} - {r.text[:200]}"
                    )
                    continue

                # Load existing document
                doc = Document(io.BytesIO(r.content))

                # Append update section with token
                doc.add_heading("Update Section", level=1)
                doc.add_paragraph(
                    f"This document has been updated. Update token: {ent['token']}\n"
                    f"Updated content to verify incremental sync functionality."
                )

                # Save updated document
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                updated_bytes = buffer.getvalue()

                # Upload updated document
                upload_url = f"/me/drive/items/{ent['id']}/content"
                r = await client.put(
                    upload_url,
                    headers={
                        "Authorization": f"Bearer {self.access_token}",
                        "Content-Type": (
                            "application/vnd.openxmlformats-officedocument."
                            "wordprocessingml.document"
                        ),
                    },
                    content=updated_bytes,
                )

                if r.status_code in (200, 201):
                    updated.append({**ent, "updated": True})
                    self.logger.info(
                        f"📝 Updated document '{ent['filename']}' with token: {ent['token']}"
                    )
                else:
                    self.logger.warning(
                        f"Failed to update document: {r.status_code} - {r.text[:200]}"
                    )

                # Brief delay between updates
                await asyncio.sleep(0.5)

        return updated

    async def delete_entities(self) -> List[str]:
        """Delete all test documents."""
        return await self.delete_specific_entities(self._test_documents)

    async def delete_specific_entities(self, entities: List[Dict[str, Any]]) -> List[str]:
        """Delete specific Word documents with retry for locked files."""
        if not entities:
            # If no specific entities provided, delete all tracked documents
            entities = self._test_documents

        if not entities:
            return []

        self.logger.info(f"🥁 Deleting {len(entities)} Word documents")
        deleted: List[str] = []

        async with httpx.AsyncClient(base_url=GRAPH, timeout=30) as client:
            # Iterate over a copy to avoid mutation issues when entities == self._test_documents
            for ent in list(entities):
                success = await self._delete_with_retry(
                    client, ent["id"], ent.get("filename", ent["id"])
                )

                if success:
                    deleted.append(ent["id"])
                    self.logger.info(f"✅ Deleted document: {ent.get('filename', ent['id'])}")

                    # Remove from tracking
                    if ent in self._test_documents:
                        self._test_documents.remove(ent)

        return deleted

    async def cleanup(self):
        """Comprehensive cleanup of all test resources."""
        self.logger.info("🧹 Starting comprehensive Word cleanup")

        cleanup_stats = {
            "documents_deleted": 0,
            "files_deleted": 0,
            "errors": 0,
        }

        try:
            async with httpx.AsyncClient(base_url=GRAPH, timeout=30) as client:
                # Delete tracked test documents
                if self._test_documents:
                    self.logger.info(
                        f"🗑️  Deleting {len(self._test_documents)} tracked documents"
                    )
                    deleted = await self.delete_specific_entities(self._test_documents.copy())
                    cleanup_stats["documents_deleted"] += len(deleted)

                # Search for and cleanup any orphaned test documents
                await self._cleanup_orphaned_files(client, cleanup_stats, "Monke_", [".docx"])

            self.logger.info(
                f"🧹 Cleanup completed: {cleanup_stats['documents_deleted']} "
                f"documents deleted, {cleanup_stats['files_deleted']} orphaned files deleted, "
                f"{cleanup_stats['errors']} errors"
            )
        except Exception as e:
            self.logger.error(f"❌ Error during comprehensive cleanup: {e}")
